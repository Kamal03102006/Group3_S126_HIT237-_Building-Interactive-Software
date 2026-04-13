"""
Generate Word documents for ADR and Testing documentation.
Run:  python generate_docs.py
Output:  ADR.docx  and  testing.docx  in the project root.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ──────────────────────────────────────────────
#  Helpers
# ──────────────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(shd)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with a dark header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "2F5496")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph("")  # spacer
    return table


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(47, 84, 150)


def body(doc, text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p


def bold_body(doc, label, text):
    p = doc.add_paragraph()
    run_b = p.add_run(label)
    run_b.bold = True
    run_b.font.size = Pt(11)
    run_n = p.add_run(text)
    run_n.font.size = Pt(11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


# ══════════════════════════════════════════════
#  ADR.docx
# ══════════════════════════════════════════════

def build_adr_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title page ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("Architecture Decision Records", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Group 3 — Indigenous Housing Management System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    date_p = doc.add_paragraph("Date: 12 April 2026")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    body(doc, "")
    body(doc, 'An Architecture Decision Record (ADR) captures an important architectural '
              'decision made along with its context and consequences.')
    body(doc, 'Template based on Michael Nygard\'s ADR format '
              '(https://github.com/joelparkerhenderson/architecture-decision-record).')

    doc.add_page_break()

    # ────────────────────────────────────────
    #  ADR-1
    # ────────────────────────────────────────
    heading(doc, "ADR-1: Standardise URL Routing and Connect All Views to Templates", level=1)
    bold_body(doc, "Date: ", "2026-04-12")

    heading(doc, "Status", level=2)
    body(doc, "Accepted")

    heading(doc, "Context", level=2)
    body(doc, 'The project root URL configuration (config/urls.py, line 5) uses '
              'include(\'housing.urls\') to delegate all housing-related routes to the housing '
              'application. However, the housing app\'s URL module was saved as url.py rather '
              'than the Django-conventional urls.py. At runtime this produced a ModuleNotFoundError '
              'because Python could not resolve housing.urls.')
    body(doc, 'Beyond the naming issue, the existing URL patterns only registered two routes — '
              'a list view and a detail view for repair requests — while the templates directory '
              'contained house_list.html and request_list.html that were not wired to any view or '
              'URL. No route existed for creating, editing, or deleting repair requests either, '
              'meaning the application only supported read-only browsing.')

    heading(doc, "Alternatives Considered", level=2)
    add_styled_table(doc,
        ["#", "Alternative", "Pros", "Cons"],
        [
            ["1", "Change the include() call to include('housing.url') "
                  "so it matches the existing filename",
             "Single one-line change; no file rename needed",
             "Violates Django convention (urls.py); confuses future "
             "contributors; every Django tutorial and linter assumes urls.py"],
            ["2", "Rename the file to urls.py and expand the URL patterns "
                  "to cover every view and template in the app",
             "Follows Django convention; every template becomes reachable; "
             "supports full CRUD workflow",
             "Requires renaming a file and adding several new URL entries"],
        ],
        col_widths=[0.3, 2.2, 2.0, 2.0],
    )

    heading(doc, "Decision", level=2)
    body(doc, 'We chose Alternative 2. The file housing/url.py was renamed to housing/urls.py, '
              'and the URL patterns were expanded to six routes:')
    add_styled_table(doc,
        ["Route", "View", "Name"],
        [
            ["/", "RepairRequestListView", "repairrequest-list"],
            ["/houses/", "DwellingListView", "dwelling-list"],
            ["/repairs/<id>/", "RepairRequestDetailView", "repairrequest-detail"],
            ["/repairs/create/", "RepairRequestCreateView", "repairrequest-create"],
            ["/repairs/<id>/edit/", "RepairRequestUpdateView", "repairrequest-update"],
            ["/repairs/<id>/delete/", "RepairRequestDeleteView", "repairrequest-delete"],
        ],
        col_widths=[1.8, 2.5, 2.2],
    )

    heading(doc, "Code Reference", level=2)
    add_styled_table(doc,
        ["File", "Lines", "What changed"],
        [
            ["config/urls.py", "5", "include('housing.urls') — unchanged, now resolves correctly"],
            ["housing/urls.py", "1–18", "Renamed from url.py; six URL patterns defined"],
        ],
        col_widths=[2.0, 1.0, 3.5],
    )

    heading(doc, "Consequences", level=2)
    bold_body(doc, "What becomes easier:", "")
    bullet(doc, "The application follows standard Django conventions, so any Django developer "
                "can navigate the codebase immediately.")
    bullet(doc, "Every template is reachable through a named URL; the {% url %} tag works in "
                "all templates without error.")
    bullet(doc, "Full CRUD operations are exposed to the browser, not just read-only views.")
    bold_body(doc, "What becomes harder or requires attention:", "")
    bullet(doc, "Any existing branch, CI script, or documentation referencing housing/url.py "
                "must be updated to housing/urls.py.")
    bullet(doc, "Adding new routes now goes through housing/urls.py, which is a longer file "
                "to maintain (18 lines vs. the original 7).")

    doc.add_page_break()

    # ────────────────────────────────────────
    #  ADR-2
    # ────────────────────────────────────────
    heading(doc, "ADR-2: Fix Template–View Mismatch Bug Discovered During Testing", level=1)
    bold_body(doc, "Date: ", "2026-04-12")

    heading(doc, "Status", level=2)
    body(doc, "Accepted")

    heading(doc, "Context", level=2)
    body(doc, 'During the first test run after cloning the repository, every page returned a '
              'Django TemplateDoesNotExist error. Investigation revealed two related bugs:')
    body(doc, '1. Views pointed to non-existent templates. RepairRequestListView set '
              'template_name = "housing/repairrequest_list.html" and RepairRequestDetailView set '
              'template_name = "housing/repairrequest_detail.html", but neither file existed. '
              'The only templates on disk were house_list.html and request_list.html.')
    body(doc, '2. Existing templates had duplicate content and wrong field names. Both '
              'house_list.html and request_list.html contained two identical <h1>Available Houses</h1> '
              'blocks and referenced model attributes house.location and house.capacity, which do '
              'not exist on the Dwelling model (housing/models.py, lines 18–42). The actual fields '
              'are house_code, address, bedrooms, and condition_status.')
    body(doc, 'These two bugs together meant the application was completely non-functional — '
              'no page could render.')

    heading(doc, "Alternatives Considered", level=2)
    add_styled_table(doc,
        ["#", "Alternative", "Pros", "Cons"],
        [
            ["1", "Rename existing templates to match the names the views expect",
             "No new files; minimal changes",
             "house_list.html is semantically a dwelling list, not a detail; "
             "leaves no template for a dwelling list view"],
            ["2", "Update the views to use the existing template filenames, fix "
                  "the template content, and create missing templates",
             "Every template matches a real view; field names align with "
             "the database schema; no duplicate HTML",
             "More files to create (3 new templates) and 2 existing templates to edit"],
        ],
        col_widths=[0.3, 2.2, 2.0, 2.0],
    )

    heading(doc, "Decision", level=2)
    body(doc, 'We chose Alternative 2:')
    bullet(doc, 'RepairRequestListView.template_name was changed to "housing/request_list.html".')
    bullet(doc, 'A new DwellingListView was created to serve "housing/house_list.html".')
    bullet(doc, 'Both existing templates were rewritten: duplicate <h1> blocks removed, '
                'field names corrected to house_code, address, community, bedrooms, condition_status.')
    bullet(doc, 'Three new templates were created: repairrequest_detail.html, '
                'repairrequest_form.html, repairrequest_confirm_delete.html.')

    heading(doc, "Code Reference", level=2)
    add_styled_table(doc,
        ["File", "Lines", "What changed"],
        [
            ["housing/views.py", "6–12", "New DwellingListView using house_list.html"],
            ["housing/views.py", "15–26", "RepairRequestListView now uses request_list.html"],
            ["housing/views.py", "29–39", "RepairRequestDetailView — prefetches updates"],
            ["housing/templates/housing/house_list.html", "1–38",
             "Rewrote with correct Dwelling model fields"],
            ["housing/templates/housing/request_list.html", "1–49",
             "Rewrote with RepairRequest model fields and navigation links"],
            ["housing/templates/housing/repairrequest_detail.html", "1–38",
             "New file — shows full request details and maintenance updates"],
            ["housing/templates/housing/repairrequest_form.html", "1–19",
             "New file — shared create/edit form with CSRF token"],
            ["housing/templates/housing/repairrequest_confirm_delete.html", "1–14",
             "New file — delete confirmation page"],
        ],
        col_widths=[2.5, 0.7, 3.3],
    )

    heading(doc, "Consequences", level=2)
    bold_body(doc, "What becomes easier:", "")
    bullet(doc, "All six URLs return HTTP 200 — the TemplateDoesNotExist error is eliminated.")
    bullet(doc, "Templates render real data from the database because they reference actual "
                "model field names.")
    bullet(doc, "Each template has a single purpose with no duplicated markup.")
    bold_body(doc, "What becomes harder or requires attention:", "")
    bullet(doc, "The original template content was fully replaced. Any in-progress front-end "
                "work must be re-applied against the new template structure.")
    bullet(doc, "Future model field changes now need updates in both the template and the view.")

    doc.add_page_break()

    # ────────────────────────────────────────
    #  ADR-3
    # ────────────────────────────────────────
    heading(doc, "ADR-3: Introduce Form Validation Through Django Generic Class-Based Views", level=1)
    bold_body(doc, "Date: ", "2026-04-12")

    heading(doc, "Status", level=2)
    body(doc, "Accepted")

    heading(doc, "Context", level=2)
    body(doc, 'The original codebase provided only two read-only views (ListView and DetailView). '
              'There was no browser-based way to create, edit, or delete a repair request — those '
              'operations required superuser access to the Django admin panel at /admin/. For a '
              'housing management system used by maintenance staff in remote communities, requiring '
              'admin credentials and navigating the admin UI is not practical.')
    body(doc, 'Django offers generic class-based views (CreateView, UpdateView, DeleteView) that '
              'automatically generate HTML forms from model field definitions and enforce server-side '
              'validation (required fields, max_length, choices, etc.). They also include built-in '
              'CSRF protection via the {% csrf_token %} template tag.')

    heading(doc, "Alternatives Considered", level=2)
    add_styled_table(doc,
        ["#", "Alternative", "Pros", "Cons"],
        [
            ["1", "Custom function-based views with manually coded form handling",
             "Full control over every validation step; no CBV learning curve",
             "Significant boilerplate per view (~30–40 lines); must manually "
             "wire CSRF, GET/POST branching, error rendering, and redirects"],
            ["2", "Django generic class-based views (CreateView, UpdateView, "
                  "DeleteView) with the fields attribute",
             "Each view is ~5 lines; validation is automatic from model fields; "
             "CSRF is built in; follows Django patterns",
             "Less fine-grained control over form layout without a custom "
             "ModelForm; complex cross-field validation requires overriding get_form()"],
        ],
        col_widths=[0.3, 2.2, 2.0, 2.0],
    )

    heading(doc, "Decision", level=2)
    body(doc, 'We chose Alternative 2. Three new views were added:')
    bullet(doc, 'RepairRequestCreateView — fields: dwelling, tenant, title, description, '
                'category, priority. On success redirects to / (repairrequest-list).')
    bullet(doc, 'RepairRequestUpdateView — fields: title, description, category, priority, '
                'status. On success redirects to /.')
    bullet(doc, 'RepairRequestDeleteView — renders a confirmation page; on POST deletes the '
                'object and redirects to /.')
    body(doc, 'A single shared template (repairrequest_form.html) handles both create and edit, '
              'toggling its heading and button text based on whether form.instance.pk exists.')

    heading(doc, "Code Reference", level=2)
    add_styled_table(doc,
        ["File", "Lines", "What changed"],
        [
            ["housing/views.py", "42–47", "RepairRequestCreateView — fields and success URL"],
            ["housing/views.py", "50–55", "RepairRequestUpdateView — fields and success URL"],
            ["housing/views.py", "58–61", "RepairRequestDeleteView — template and success URL"],
            ["housing/urls.py", "14", "repairs/create/ route"],
            ["housing/urls.py", "15", "repairs/<int:pk>/edit/ route"],
            ["housing/urls.py", "16", "repairs/<int:pk>/delete/ route"],
            ["housing/templates/housing/repairrequest_form.html", "10–17",
             "<form method=\"post\"> with {% csrf_token %} and {{ form.as_table }}"],
            ["housing/models.py", "100–101",
             "title is CharField(max_length=150) — validation enforced automatically"],
            ["housing/models.py", "102",
             "description is TextField() — required by default"],
        ],
        col_widths=[2.5, 0.7, 3.3],
    )

    heading(doc, "Consequences", level=2)
    bold_body(doc, "What becomes easier:", "")
    bullet(doc, "Maintenance staff can create, edit, and delete repair requests through the "
                "browser without admin credentials.")
    bullet(doc, "Server-side validation is automatic: empty required fields, values outside "
                "choices, and strings exceeding max_length are all rejected with clear error messages.")
    bullet(doc, "CSRF protection is enforced on every POST, protecting against cross-site "
                "request forgery attacks.")
    bullet(doc, "Code is minimal — three views totalling ~20 lines, easy to read and maintain.")
    bold_body(doc, "What becomes harder or requires attention:", "")
    bullet(doc, "The auto-generated form uses HTML <table> layout. A more polished UI will "
                "require a custom ModelForm with widget overrides or a CSS/JS framework.")
    bullet(doc, "If cross-field validation is needed later, a custom clean() method must be "
                "added to a ModelForm subclass.")

    doc.save("ADR.docx")
    print("Created: ADR.docx")


# ══════════════════════════════════════════════
#  testing.docx
# ══════════════════════════════════════════════

TEST_CASES = [
    {
        "id": "TC-01",
        "name": "Home Page (Repair Request List) Loads Successfully",
        "input": "Navigate to http://127.0.0.1:8000/",
        "expected": 'Page displays "Repair Requests" heading with a table of repair requests '
                    '(or an empty-state message). Navigation links to Houses and '
                    '"Submit New Request" are visible.',
        "actual": "Page loads with status 200. Heading, table, and navigation links render correctly.",
        "status": "Pass",
    },
    {
        "id": "TC-02",
        "name": "Dwelling List Page Loads Successfully",
        "input": "Navigate to http://127.0.0.1:8000/houses/",
        "expected": 'Page displays "Available Houses" heading with a table showing house code, '
                    "address, community, bedrooms, and condition.",
        "actual": "Page loads with status 200. Table columns render correctly. "
                  "Empty-state message shown when no data exists.",
        "status": "Pass",
    },
    {
        "id": "TC-03",
        "name": "Repair Request Detail Page",
        "input": "Create a RepairRequest via Django admin, then navigate to "
                 "http://127.0.0.1:8000/repairs/1/",
        "expected": "Page displays the repair request title, dwelling, tenant, category, "
                    "priority, status, description, and timestamps. Maintenance updates section is shown.",
        "actual": "Detail page renders all fields correctly with status 200.",
        "status": "Pass",
    },
    {
        "id": "TC-04",
        "name": "Repair Request Detail — Non-Existent ID Returns 404",
        "input": "Navigate to http://127.0.0.1:8000/repairs/9999/",
        "expected": "Server responds with HTTP 404 Not Found.",
        "actual": "Django returns a 404 error page.",
        "status": "Pass",
    },
    {
        "id": "TC-05",
        "name": "Create Repair Request — Valid Form Submission",
        "input": "Navigate to http://127.0.0.1:8000/repairs/create/. "
                 "Fill all required fields (dwelling, title, description, category, priority) and submit.",
        "expected": "Form submits successfully. User is redirected to the repair request list page. "
                    "New request appears in the table.",
        "actual": "POST succeeds with 302 redirect to /. New repair request visible in list.",
        "status": "Pass",
    },
    {
        "id": "TC-06",
        "name": "Create Repair Request — Empty Required Fields (Form Validation)",
        "input": "Navigate to http://127.0.0.1:8000/repairs/create/. "
                 "Leave the title and description fields empty and submit.",
        "expected": "Form does not submit. Validation errors are displayed for required fields "
                    "(title, description, dwelling, category).",
        "actual": 'Form re-renders with validation errors: "This field is required" shown '
                  "next to each empty required field.",
        "status": "Pass",
    },
    {
        "id": "TC-07",
        "name": "Update Repair Request — Change Status",
        "input": "Navigate to http://127.0.0.1:8000/repairs/1/edit/. "
                 'Change the status field from "Reported" to "In Progress" and submit.',
        "expected": "Form submits successfully. User is redirected to the list page. "
                    'The request\'s status is updated to "In Progress".',
        "actual": "POST succeeds with 302 redirect. Status updated in both the list and detail views.",
        "status": "Pass",
    },
    {
        "id": "TC-08",
        "name": "Update Repair Request — Form Validation on Edit",
        "input": "Navigate to http://127.0.0.1:8000/repairs/1/edit/. "
                 "Clear the title field and submit.",
        "expected": "Form does not submit. Validation error shown for the title field.",
        "actual": 'Form re-renders with "This field is required" error on the title field.',
        "status": "Pass",
    },
    {
        "id": "TC-09",
        "name": "Delete Repair Request — Confirmation and Deletion",
        "input": "Navigate to http://127.0.0.1:8000/repairs/1/delete/. "
                 'Confirm deletion by clicking "Yes, delete".',
        "expected": "Request is deleted. User is redirected to the list page. "
                    "The deleted request no longer appears.",
        "actual": "POST succeeds with 302 redirect. Request removed from list.",
        "status": "Pass",
    },
    {
        "id": "TC-10",
        "name": "Delete Repair Request — Cancel Returns to Detail",
        "input": "Navigate to http://127.0.0.1:8000/repairs/1/delete/. "
                 'Click "Cancel" instead of confirming.',
        "expected": "User is taken back to the repair request detail page. "
                    "The request is not deleted.",
        "actual": "Cancel link navigates to detail page. Request still exists.",
        "status": "Pass",
    },
    {
        "id": "TC-11",
        "name": "Navigation — Links Between Pages Work Correctly",
        "input": 'From the repair request list page, click "View Houses". '
                 'From the houses page, click "View Repair Requests". '
                 'From the list, click a request title to view details. '
                 'From details, click "Back to Requests".',
        "expected": "All navigation links resolve correctly and load the expected pages "
                    "without 404 errors.",
        "actual": "All links work. Pages load with HTTP 200.",
        "status": "Pass",
    },
    {
        "id": "TC-12",
        "name": "Admin Panel — Model Registration",
        "input": "Navigate to http://127.0.0.1:8000/admin/ and log in with superuser credentials.",
        "expected": "Admin panel shows registered models: Community, Dwelling, Tenant, "
                    "RepairRequest, MaintenanceUpdate. Each model has list display, search, "
                    "and filter configurations.",
        "actual": "All five models appear in admin. List display columns, search fields, "
                  "and filters work as configured in admin.py.",
        "status": "Pass",
    },
    {
        "id": "TC-13",
        "name": "URL Routing — Invalid URL Returns 404",
        "input": "Navigate to http://127.0.0.1:8000/nonexistent-page/",
        "expected": "Server responds with HTTP 404 Not Found.",
        "actual": "Django returns 404 error page.",
        "status": "Pass",
    },
    {
        "id": "TC-14",
        "name": "CSRF Protection on Forms",
        "input": "Attempt a POST request to http://127.0.0.1:8000/repairs/create/ "
                 "without a CSRF token (e.g., via curl).",
        "expected": "Server responds with HTTP 403 Forbidden.",
        "actual": 'Django returns 403 Forbidden with "CSRF verification failed" message.',
        "status": "Pass",
    },
]


def build_testing_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title page ──
    doc.add_paragraph("")
    doc.add_paragraph("")
    title = doc.add_heading("Testing Document", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Group 3 — Indigenous Housing Management System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    date_p = doc.add_paragraph("Date: 12 April 2026")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Test environment ──
    heading(doc, "Test Environment", level=1)
    add_styled_table(doc,
        ["Item", "Details"],
        [
            ["Framework", "Django 4.2.29"],
            ["Database", "SQLite 3 (development)"],
            ["Python", "3.14"],
            ["Server", "Django development server (manage.py runserver)"],
        ],
        col_widths=[2.0, 4.5],
    )

    # ── Test cases ──
    heading(doc, "Test Cases", level=1)

    for tc in TEST_CASES:
        heading(doc, f'{tc["id"]}: {tc["name"]}', level=2)
        add_styled_table(doc,
            ["Field", "Details"],
            [
                ["Input", tc["input"]],
                ["Expected Output", tc["expected"]],
                ["Actual Output", tc["actual"]],
                ["Status", tc["status"]],
            ],
            col_widths=[1.5, 5.0],
        )

    # ── Summary ──
    doc.add_page_break()
    heading(doc, "Summary", level=1)

    passed = sum(1 for tc in TEST_CASES if tc["status"] == "Pass")
    failed = len(TEST_CASES) - passed

    add_styled_table(doc,
        ["Total Tests", "Passed", "Failed"],
        [
            [str(len(TEST_CASES)), str(passed), str(failed)],
        ],
        col_widths=[2.2, 2.2, 2.2],
    )

    body(doc, "All pages load successfully, form validation works on both create and update "
              "views, CRUD operations function correctly, navigation links are properly "
              "connected, and Django's CSRF protection is enforced.")

    doc.save("testing.docx")
    print("Created: testing.docx")


# ══════════════════════════════════════════════
#  Main
# ══════════════════════════════════════════════

if __name__ == "__main__":
    build_adr_doc()
    build_testing_doc()
    print("Done — both Word documents generated.")
